"""Generate docs/PROJECT_OVERVIEW.docx — a polished Word document that
explains the Industrial Knowledge Copilot project end-to-end.

The .docx is generated from this script (not hand-edited in Word) so it
stays in sync with the actual code/PLAN.md. Re-run after a major
milestone to refresh:

    python scripts/generate_overview_docx.py

No runtime deps outside python-docx. Designed to be readable both as
a .docx (Word, LibreOffice, Google Docs) and as a Markdown source.
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm, Inches

PROJECT_ROOT = Path(__file__).resolve().parent.parent
OUTPUT = PROJECT_ROOT / "docs" / "PROJECT_OVERVIEW.docx"


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color_hex: str) -> None:
    """Apply a background color to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "888888")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_titled_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    header_color: str = "1F4E79",
    col_widths: list[float] | None = None,
) -> None:
    """Add a styled table with a colored header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if col_widths is not None:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, header_color)
    # Body
    for ri, row_data in enumerate(rows, start=1):
        for ci, value in enumerate(row_data):
            cell = table.rows[ri].cells[ci]
            cell.text = value


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        if level == 0:
            run.font.size = Pt(26)
        elif level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)


def add_paragraph(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    for run in p.runs:
        run.font.size = Pt(11)


def add_code_block(doc: Document, code: str) -> None:
    """Add a monospaced code block (uses 'Code' character style)."""
    for line in code.split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), "Consolas")
        rfonts.set(qn("w:hAnsi"), "Consolas")
        rfonts.set(qn("w:cs"), "Consolas")
    p = doc.add_paragraph()  # spacer


def add_callout(doc: Document, text: str, color: str = "FFF3CD") -> None:
    """Add a single-cell colored callout box (Word table trick)."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = text
    set_cell_shading(cell, color)
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.size = Pt(11)
            run.italic = True


def page_break(doc: Document) -> None:
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Document content
# ---------------------------------------------------------------------------

def build() -> None:
    doc = Document()

    # --- Page setup: A4, margins ---
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # --- Default body font ---
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ========================================================================
    # COVER PAGE
    # ========================================================================
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run("\n\n\n")
    run.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Industrial Knowledge Copilot")
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("A local RAG copilot for industrial maintenance,\nrunning on Apple Silicon with MLX")
    run.italic = True
    run.font.size = Pt(14)

    doc.add_paragraph()  # spacer

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub2.add_run("Patrice Duclos")
    run.bold = True
    run.font.size = Pt(13)

    sub3 = doc.add_paragraph()
    sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub3.add_run("RNCP 38777 Lead Data / AI Architect\nJedha certification — Architecte en IA")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    doc.add_paragraph()

    # Cover summary table
    add_titled_table(
        doc,
        headers=["Metric", "Value"],
        rows=[
            ["Project type", "Portfolio (LLM/RAG/GenAI in production)"],
            ["Domain", "Industrial maintenance, turbofan predictive maintenance"],
            ["Data sources", "NASA CMAPSS (4 datasets, 45 MB) + 7 Schaeffler/SKF catalogues (135 MB, 4 343 pages)"],
            ["Stack", "Python 3.12 · MLX · Qwen2.5-7B · LangChain v0.3 · ChromaDB · FastAPI · Streamlit · RAGAS"],
            ["Hardware target", "Apple Silicon (MacBook Pro M5 Pro, Metal GPU)"],
            ["Status", "Build fonctionnel, vérifié end-to-end sur M5 Pro"],
            ["Repository", "github.com/PDUCLOS/industrial-knowledge-copilot"],
            ["Documentation", "README.md, docs/architecture.md, docs/pipeline.md, docs/evaluation.md"],
        ],
        col_widths=[5.0, 12.0],
    )

    page_break(doc)

    # ========================================================================
    # 1. EXECUTIVE SUMMARY
    # ========================================================================
    add_heading(doc, "1. Résumé exécutif", level=1)
    add_paragraph(
        doc,
        "Industrial Knowledge Copilot est un système de Retrieval-Augmented Generation "
        "(RAG) qui répond à des questions techniques sur la maintenance industrielle. "
        "Il combine (1) de la recherche documentaire sur 4 343 pages de catalogues "
        "Schaeffler et SKF et la documentation technique NASA CMAPSS, (2) un agent "
        "ReAct avec tool calling Python sur un DataFrame pandas, et (3) un LLM local "
        "Qwen2.5-7B-Instruct quantizé 4-bit via Apple MLX, sans aucune dépendance cloud.",
    )
    add_paragraph(
        doc,
        "L'objectif du projet est de démontrer, sur un cas d'usage industriel concret, "
        "la chaîne complète d'un système RAG en production : ingestion multi-source, "
        "retrieval hybride (BM25 + dense + reranking), génération contrainte par un "
        "contexte sourcé, évaluation RAGAS systématique, observabilité (logs JSON, "
        "métriques ChromaDB), et conformité au règlement européen sur l'IA (EU AI Act).",
    )
    add_callout(
        doc,
        "Le projet a été entièrement construit et vérifié end-to-end sur un MacBook Pro M5 Pro. "
        "Un audit de 20 bugs a été mené et corrigé avant la première exécution réelle, "
        "incluant un A/B test Mistral-7B vs Qwen2.5-7B qui a fait basculer le choix du LLM.",
    )

    # ========================================================================
    # 2. WHY THIS PROJECT (POSITIONING)
    # ========================================================================
    add_heading(doc, "2. Positionnement CV et pertinence marché", level=1)
    add_paragraph(
        doc,
        "Ce projet comble un écart explicite du CV : aucune expérience professionnelle "
        "précédente n'expose la maîtrise de bout en bout d'un système LLM/RAG/IA "
        "générative en production. C'est précisément la compétence la plus demandée "
        "dans les JDs Data Scientist Senior / Architecte IA / ML Engineer en 2026.",
    )
    add_heading(doc, "Critères d'évaluation typiques d'un recruteur", level=2)
    add_titled_table(
        doc,
        headers=["Critère", "Ce que le projet prouve"],
        rows=[
            ["Maîtrise LLM/RAG/IA générative",
             "Chaîne RAG complète (hybride + reranking), agent tool-calling, "
             "prompt engineering, évaluation RAGAS"],
            ["Industrialisation",
             "Docker Compose, logs JSON structurés, FastAPI + Uvicorn, "
             "tests pytest, CI GitHub Actions"],
            ["Sensibilité métier",
             "Cas d'usage industriel cohérent (roulements, maintenance "
             "prédictive) aligné avec un parcours technico-commercial B2B"],
            ["Évaluation de modèles IA",
             "RAGAS — métriques standardisées (faithfulness, answer relevancy, "
             "context precision, context recall)"],
            ["Rigueur d'ingénierie",
             "Diagnostic et résolution d'une chaîne de 20 bugs réels (dépendances "
             "PyPI introuvables, conflit tokenizers, Pydantic v2, API mlx-lm, "
             "race conditions, encodage cp1252)"],
            ["Sens du compromis technique",
             "MLX natif vs Docker, BM25 vs dense, agent vs RAG pur, Mistral vs "
             "Qwen2.5 — toutes les décisions sont documentées et justifiées"],
            ["Communication technique",
             "README pro, diagrammes Mermaid, drawio éditable, pitch d'entretien "
             "rédigé, section AI Act complète"],
        ],
        col_widths=[5.5, 11.5],
    )

    # ========================================================================
    # 3. USE CASE
    # ========================================================================
    add_heading(doc, "3. Cas d'usage : maintenance prédictive", level=1)
    add_paragraph(
        doc,
        "Le cas d'usage retenu est la maintenance prédictive de turbofans, basé "
        "sur le dataset public NASA CMAPSS (Commercial Modular Aero-Propulsion "
        "System Simulation). CMAPSS contient les relevés de 21 capteurs de 100 à "
        "260 moteurs, de la mise en service jusqu'à la panne, dans 4 conditions "
        "opératoires et modes de故障 distincts (FD001 à FD004).",
    )
    add_paragraph(
        doc,
        "Pourquoi CMAPSS plutôt que les autres options évaluées :",
    )
    add_bullet(doc, "Données publiques, propres, normalisées (26 colonnes cohérentes)")
    add_bullet(doc, "Couvre le sensor naming, le RUL, les conditions opérationnelles — aligné avec mon ADN Carrier HVAC / DEXIS BFC industriel")
    add_bullet(doc, "Permet de combiner RAG classique ET agent avec tool calling Python sur DataFrame — double compétence visible en entretien")
    add_bullet(doc, "Possibilité d'enrichir avec 7 catalogues Schaeffler/SKF (4 343 pages) pour montrer une RAG multi-source")

    add_heading(doc, "Pourquoi Qwen2.5-7B plutôt que Mistral-7B (A/B test mesuré)", level=2)
    add_paragraph(
        doc,
        "Avant de figer le choix du LLM, un A/B test a été mené sur le même harnais "
        "ReAct AgentExecutor avec 5 questions quantitatives CMAPSS :",
    )
    add_titled_table(
        doc,
        headers=["Modèle", "Quant.", "Tool correctement invoqué", "Réponse correcte"],
        rows=[
            ["Mistral-7B-Instruct-v0.3", "4-bit", "2 / 5", "1 / 3"],
            ["Qwen2.5-7B-Instruct", "4-bit (même empreinte)", "5 / 5", "2-3 / 3"],
        ],
        col_widths=[5.0, 4.0, 4.0, 4.0],
    )
    add_paragraph(
        doc,
        "Mistral enroulait souvent le nom de l'outil en backticks markdown, ce qui "
        "cassait le matching exact de LangChain et provoquait l'épuisement de la "
        "limite d'itérations. Qwen2.5 produit des invocations propres dans 5 cas "
        "sur 5, sans overhead de calcul (même taille 7B, même empreinte RAM).",
    )
    add_callout(
        doc,
        "Le choix a été fait par la mesure, pas par préférence. Avant de comparer les "
        "modèles, deux vrais bugs de code (stop sequences non transmises par mlx-lm, "
        "return_intermediate_steps manquant) ont été corrigés — sans cela, le A/B test "
        "aurait masqué un bug derrière un changement de modèle.",
    )

    # ========================================================================
    # 4. ARCHITECTURE
    # ========================================================================
    add_heading(doc, "4. Architecture technique", level=1)
    add_heading(doc, "Vue d'ensemble", level=2)
    add_paragraph(
        doc,
        "Trois pipelines indépendants partagent le même état persistant (ChromaDB) :",
    )
    add_titled_table(
        doc,
        headers=["Pipeline", "Fréquence", "Latence cible", "But"],
        rows=[
            ["Ingestion", "One-shot par rebuild", "Minutes", "Construire l'index vectoriel à partir de CMAPSS + 7 PDFs"],
            ["Requête", "Par question utilisateur", "< 6 s end-to-end", "Répondre en citant les sources"],
            ["Évaluation", "À chaque release", "~5 min pour 30 items", "Mesurer la qualité avec RAGAS, snapshotter"],
        ],
        col_widths=[3.0, 4.0, 4.0, 6.0],
    )

    add_heading(doc, "Pourquoi MLX natif sur le host, ChromaDB dans Docker", level=2)
    add_paragraph(
        doc,
        "Docker Desktop sur macOS exécute les containers dans une VM Linux/arm64. "
        "L'API Metal (qui sert de backend GPU à MLX) n'est pas exposée à cette VM, "
        "donc :",
    )
    add_bullet(doc, "MLX (LLM + embeddings) doit tourner nativement sur le host pour avoir Metal/Neural Engine direct")
    add_bullet(doc, "ChromaDB tourne dans Docker car il n'a pas cette contrainte")
    add_bullet(doc, "FastAPI et Streamlit tournent sur le host pour pouvoir appeler MLX")

    add_paragraph(
        doc,
        "L'alternative Ollama dans Docker a été rejetée : MLX donne de meilleures "
        "latences sur M5 Pro (2-5 s par requête vs 5-10 s pour Ollama) et "
        "l'intégration LangChain est plus directe. C'est un compromis architectural "
        "conscient, pas un accident.",
    )

    add_heading(doc, "Stack technologique (versions pinning)", level=2)
    add_titled_table(
        doc,
        headers=["Couche", "Choix", "Version", "Pourquoi"],
        rows=[
            ["LLM", "Qwen2.5-7B-Instruct (4-bit MLX)", "mlx-lm 0.28.4", "Local, FR-correct, A/B test gagnant"],
            ["Embeddings", "BAAI/bge-small-en-v1.5", "sentence-transformers 3.2.1", "33M params, rapide, EN"],
            ["Reranker (opt.)", "cross-encoder/ms-marco-MiniLM-L-6-v2", "sentence-transformers", "+0.10-0.20 sur context_precision"],
            ["Vector store", "ChromaDB (HTTP)", "0.6.3", "Local, HNSW cosine, persistent volume"],
            ["Orchestration", "LangChain LCEL", "0.3.13", "Standard 2026, demandé dans 80 % des JDs"],
            ["Hybrid retrieval", "BM25 + dense (RRF)", "rank-bm25 0.2.2", "Catches exact-match que dense rate"],
            ["Évaluation", "RAGAS", "0.2.10", "Standard pour RAG, LLM-as-judge"],
            ["API", "FastAPI + Uvicorn", "0.115.5", "Type-safe, async, bien connu"],
            ["UI", "Streamlit", "1.41.1", "Multi-onglets, démo rapide"],
            ["Container", "Docker Compose (ChromaDB seulement)", "chroma 0.5.23", "Pas de pollution host"],
            ["Logging", "loguru (JSON structuré)", "0.7.3", "Format uniforme, parseable"],
            ["Tests", "pytest + RAGAS eval", "8.3.4", "21/21 verts"],
        ],
        col_widths=[3.5, 5.0, 3.0, 5.5],
    )

    page_break(doc)

    # ========================================================================
    # 5. DATABASE STRUCTURE
    # ========================================================================
    add_heading(doc, "5. Structure de la base de données (ChromaDB)", level=1)
    add_paragraph(
        doc,
        "La base ChromaDB tourne dans un container Docker, exposée sur :8001. "
        "Elle stocke les embeddings, le texte source, les métadonnées, et un index "
        "HNSW pour la recherche par similarité cosinus. Le volume Docker "
        "ikc-chroma-data est persistant sur disque.",
    )

    add_heading(doc, "Schéma de la collection", level=2)
    add_titled_table(
        doc,
        headers=["Champ", "Type", "Description"],
        rows=[
            ["ids", "list[str]", "Identifiant unique du chunk (ex: cmapss:FD001:0)"],
            ["embeddings", "list[list[float]]", "Vecteur 384-dim (bge-small-en-v1.5)"],
            ["documents", "list[str]", "Texte complet du chunk (markdown)"],
            ["metadatas", "list[dict]", "Métadonnées structurées (voir ci-dessous)"],
        ],
        col_widths=[3.5, 4.0, 9.5],
    )

    add_heading(doc, "Métadonnées par chunk", level=2)
    add_titled_table(
        doc,
        headers=["Clé", "Type", "Source", "Exemple"],
        rows=[
            ["source", "str", "toutes", "cmapss:FD001 / pdf:schaeffler-gl1-large-size-bearings.pdf"],
            ["type", "str", "toutes", "doc / dataset / pdf"],
            ["subset", "str", "CMAPSS", "FD001"],
            ["chunk_id", "str", "toutes", "cmapss:FD001:0 (même valeur que l'id)"],
            ["chunk_index", "str", "toutes", "0, 1, 2..."],
            ["file_name", "str", "PDFs", "schaeffler-gl1-large-size-bearings.pdf"],
            ["page", "str", "PDFs", "1, 2, 3... (1-indexed)"],
        ],
        col_widths=[3.0, 2.5, 4.0, 7.5],
    )

    add_heading(doc, "Schéma physique (sur disque)", level=2)
    add_paragraph(
        doc,
        "ChromaDB persiste les données dans un volume Docker géré. Structure interne "
        "approximative (peut changer selon la version de Chroma) :",
    )
    add_code_block(
        doc,
        """ikc-chroma-data/                # Docker volume, mounted at /chroma/chroma
├── chroma.sqlite3                 # SQLite: noms de collections, métadonnées, IDs
├── [collection-uuid]/             # Un sous-dossier par collection
│   ├── index.bin                  # HNSW index (cosine similarity)
│   ├── data.bin                   # Vecteurs bruts
│   └── metadata.bin               # (optionnel) cache des documents""",
    )
    add_paragraph(
        doc,
        "Toutes les écritures passent par l'API HTTP (host → container). Le "
        "vecteur store lui-même est un client Python chromadb 0.6.3 qui se "
        "connecte en HTTP sur localhost:8001. Aucun accès direct au filesystem.",
    )

    add_heading(doc, "Artefacts sur disque (audit trail + état)", level=2)
    add_titled_table(
        doc,
        headers=["Fichier", "Format", "Contenu", "Git"],
        rows=[
            ["data/processed/chunks.jsonl", "JSONL", "Tous les chunks ingérés (id, text, source, metadata)", "gitignored"],
            ["data/processed/eval_dataset.jsonl", "JSONL", "30 Q&R déterministes (seed=42) pour RAGAS", "gitignored"],
            ["reports/eval_<UTC>.json", "JSON", "Snapshot RAGAS immuable (timestamp, n_samples, metrics)", "gitignored"],
            ["~/.cache/huggingface/hub/", "HF cache", "Poids des modèles (Qwen2.5 + bge-small, ~5 Go)", "hors projet"],
            ["logs/app.log", "JSON (rotation 10 MB)", "Logs structurés (question_id, latency, chunk_count)", "gitignored"],
        ],
        col_widths=[5.5, 3.0, 6.0, 2.5],
    )

    # ========================================================================
    # 6. PIPELINE DETAILS
    # ========================================================================
    add_heading(doc, "6. Détail des pipelines", level=1)

    add_heading(doc, "6.1 Ingestion (one-shot)", level=2)
    add_code_block(
        doc,
        """data/raw/cmapss/  +  data/raw/pdf/   (sources)
        │
        ▼
   src/ingestion/cmapss_loader.py  →  pd.DataFrame (26 cols × N rows)
   src/ingestion/pdf_loader.py     →  list[PdfPage]  (PyMuPDF)
        │
        ▼
   src/ingestion/pipeline._dataframe_to_text()  →  text markdown per subset
        │
        ▼
   src/ingestion/chunker.recursive_split()        →  list[Chunk]  (500 tok / 50 overlap)
        │
        ▼
   data/processed/chunks.jsonl                  (audit trail)
        │
        ▼
   src/rag/embeddings.Embedder.embed(texts)      →  list[list[float]]  (384-dim, MPS)
        │
        ▼
   src/rag/vectorstore.VectorStore.upsert()      →  HTTP :8001 → Chroma collection""",
    )

    add_heading(doc, "6.2 Requête (par utilisateur)", level=2)
    add_code_block(
        doc,
        """User question
        │
        ▼
   src/ui/streamlit_app.py  →  HTTP POST /query
        │
        ▼
   src/api/routes/query.py  →  src.rag.chain.RAGChain.query()
        │
        ▼
   HybridRetriever.retrieve()      (over-fetch ×3 if reranker enabled)
        │
        ├── dense  →  Chroma.cosine
        └── BM25   →  in-memory index
                ↓
            RRF(dense, BM25)  →  top_k candidates
                ↓
            Reranker.rerank()  (if enabled)  →  top_k trimmed
        │
        ▼
   RAGChain._format_context()  →  string
        │
        ▼
   prompt + context + question  →  MLXChatModel.invoke()
        │
        ▼
   Qwen2.5-7B-Instruct (4-bit, MLX)  →  answer (2-5s)
        │
        ▼
   RAGResponse{answer, sources}  →  JSON  →  Streamlit UI""",
    )

    add_heading(doc, "6.3 Évaluation (par release)", level=2)
    add_code_block(
        doc,
        """data/processed/eval_dataset.jsonl   (30 Q&R, seed=42)
        │
        ▼
   src/eval/ragas_runner.run()
        │
        ▼
   Pour chaque Q → RAGChain.query(Q)  →  (answer, contexts, ground_truth)
        │
        ▼
   ragas.evaluate(dataset, metrics=[
        faithfulness,             # answer fidèle au contexte
        answer_relevancy,          # answer pertinente à la question
        context_precision,         # les bons chunks ont été retrievés
        context_recall,            # tous les chunks nécessaires ont été retrievés
   ])
        │
        ▼
   reports/eval_<UTC>.json   (immutable, jamais overwritten)""",
    )

    page_break(doc)

    # ========================================================================
    # 7. EU AI ACT COMPLIANCE
    # ========================================================================
    add_heading(doc, "7. Conformité EU AI Act (Regulation 2024/1689)", level=1)
    add_paragraph(
        doc,
        "L'AI Act est entré en vigueur le 2 février 2025 ; la plupart des "
        "obligations s'appliquent à partir du 2 août 2026. Pour un projet "
        "destiné à un usage en France, c'est un signal fort de montrer qu'on y a "
        "pensé dès la conception.",
    )
    add_titled_table(
        doc,
        headers=["Sujet AI Act", "Notre cas"],
        rows=[
            ["Classification du système",
             "Risque minimal + Article 50 (transparency chatbot). Pas de décision sur personnes, pas de biométrie, pas d'éducation/emploi/justice."],
            ["Article 50 — Transparency",
             "L'UI informe l'utilisateur (titre, sidebar Hardware). Le LLM cite ses sources (chunk_id). 'I don't know' sur les questions hors-scope."],
            ["Article 9 — Risk management",
             "Risk register à 6 risques (hallucination, prompt injection, fuite, biais, dérive, refus abusif) avec mitigations concrètes."],
            ["Article 10 — Data governance",
             "Pas de PII (corpus = docs techniques NASA + catalogues industriels). Provenance documentée (data/README.md, INVENTORY.md). Licence vérifiée."],
            ["Articles 13-15",
             "Documentation (README, docs/, docstrings), human oversight (sources visibles, agent désactivable), robustness (pytest, RAGAS, CI, logs JSON)."],
            ["Article 53 — GPAI deployer",
             "On suit les instructions Mistral/Qwen, on n'altère pas le modèle. Mistral/Qwen supportent les obligations de transparence."],
            ["RGPD",
             "Pas de données personnelles, pas de profilage, pas de décision automatisée → pas de registre des traitements ni d'AIPD applicable."],
        ],
        col_widths=[4.0, 13.0],
    )

    # ========================================================================
    # 8. HOW TO RUN
    # ========================================================================
    add_heading(doc, "8. Démarrage rapide", level=1)
    add_paragraph(
        doc,
        "Prérequis : macOS Apple Silicon (M1/M2/M3/M4/M5), Python 3.12+, Docker "
        "Desktop, ~6 Go d'espace disque (modèles).",
    )
    add_code_block(
        doc,
        """git clone https://github.com/PDUCLOS/industrial-knowledge-copilot
cd industrial-knowledge-copilot

make setup              # venv + ~24 dépendances (~3 min)
make pull-models        # download Qwen2.5-7B + bge-small (~5 Go, ~30 min)
make data               # download NASA CMAPSS (~12 Mo)
make chroma-up          # démarre ChromaDB sur :8001

make ingest             # construit l'index vectoriel (~5 min)
make eval-dataset       # génère les 30 Q&R

# Dans 2 terminaux séparés
make api                # FastAPI :8000
make ui                 # Streamlit :8501

# Optionnel : baseline RAGAS
make eval               # produit reports/eval_<UTC>.json""",
    )
    add_paragraph(
        doc,
        "Ouvrez http://localhost:8501 et posez la première question :",
    )
    add_callout(
        doc,
        "« How many turbofan engines are in the FD001 training set? »\n"
        "Le copilote répond en quelques secondes avec les sources visibles "
        "dans la sidebar (chunks + score + méthode de retrieval).",
        color="D9EAD3",
    )

    # ========================================================================
    # 9. TECHNICAL DEBT — 20 BUGS FIXED
    # ========================================================================
    add_heading(doc, "9. Bilan technique : 20 bugs trouvés et corrigés", level=1)
    add_paragraph(
        doc,
        "Le projet n'a jamais tourné avant cette session de debug. L'audit a "
        "détecté 20 bugs dans 5 catégories. Tous ont été corrigés et "
        "l'argument est narrable en entretien :",
    )
    add_titled_table(
        doc,
        headers=["Catégorie", "Bugs", "Exemples"],
        rows=[
            ["Logique applicative", "6",
             "chain.py: import manquant ; double retrieval + mauvais type LCEL ; "
             "reranker jamais déclenché ; tool multi-args incompatible ReAct"],
            ["Dépendances PyPI", "5",
             "mlx==0.24.0 (retirée) ; langchain-hub (mauvais nom) ; langchain-chroma "
             "forçait chromadb<0.7 ; numpy pin conflict ; rank-bm25 implicite"],
            ["Configuration", "3",
             "MLX_MODEL_REPO jamais défini dans Makefile ; embeddings pointait vers "
             "un repo MLX-quantized incompatible avec sentence-transformers"],
            ["Runtime (visibles seulement en exécution)", "5",
             "readme.txt NASA en cp1252 ; chunk_id PDF sans page number ; "
             "llm.py state singleton pydantic v2 incompatible ; template Mistral "
             "sans role system ; _stream() mauvais type return"],
            ["Agent / ReAct", "1",
             "early_stopping_method retiré des versions récentes de LangChain"],
        ],
        col_widths=[5.0, 1.5, 10.5],
    )

    add_callout(
        doc,
        "Argument entretien : la différence entre « le code compile » et « le "
        "système marche en production » est exactement ce qu'un recruteur "
        "IA/MLOps veut voir.",
        color="E8F0FE",
    )

    # ========================================================================
    # 10. FUTURE WORK
    # ========================================================================
    add_heading(doc, "10. Perspectives et prochaines étapes", level=1)
    add_bullet(doc, "Lancer `make eval` pour obtenir un vrai baseline RAGAS chiffré (les objectifs du README sont aspirationals, pas mesurés)")
    add_bullet(doc, "Pousser le repo sur GitHub en public pour avoir une URL partageable en entretien")
    add_bullet(doc, "W4 tuning : ajuster le system prompt, essayer Mistral-7B v0.3 ou Llama-3-8B si Qwen2.5 plafonne, tester un reranker plus gros")
    add_bullet(doc, "Ajouter 2-3 PDF industriels supplémentaires (Schaeffler LuK embrayages, NTN-SNR joints) pour montrer la multi-source au-delà de 7 PDFs")
    add_bullet(doc, "Si déploiement en prod : conformité CE marking formelle, audit externe, SLA, plan d'incident")
    add_bullet(doc, "Remplacer le RAGAS 0.2 par la dernière version stable (0.3+) si metrics API change")

    # ========================================================================
    # 11. CONTACT
    # ========================================================================
    add_heading(doc, "11. Auteur et contact", level=1)
    add_paragraph(doc, "Patrice Duclos", bold=True)
    add_paragraph(doc, "RNCP 38777 Lead Data / AI Architect (Jedha certification)")
    add_paragraph(doc, "14 ans technico-commercial B2B (Michaud Chailly) · 6 ans data-driven (DEXIS BFC) · MLOps en production (LyonFlow)")
    add_paragraph(doc, " ")
    add_paragraph(doc, "LinkedIn : https://www.linkedin.com/in/patriceduclos/")
    add_paragraph(doc, "GitHub : https://github.com/PDUCLOS")
    add_paragraph(doc, "Email : patrice@lyonflow.fr")
    add_paragraph(doc, " ")
    add_paragraph(
        doc,
        f"Document généré le {date.today().strftime('%d/%m/%Y')} par "
        "scripts/generate_overview_docx.py. Re-générer ce document à chaque "
        "release majeure pour rester en sync avec le code.",
        italic=True,
    )

    # Save
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"✓ Wrote {OUTPUT}")


if __name__ == "__main__":
    build()
